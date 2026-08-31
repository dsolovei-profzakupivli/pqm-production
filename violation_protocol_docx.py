"""DOCX protocol generator for customer violation reports.

The module edits copies of the three approved Word packages.  It deliberately
contains no database or Prozorro access: readiness and freshness are enforced
by the HTTP/service layer.
"""
from __future__ import annotations

import re
import shutil
import os
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PACKAGED_TEMPLATE_DIR = Path(__file__).with_name("templates") / "violation_protocols"
TEMPLATE_DIR = (Path(os.environ["PQM_DATA_DIR"]) / "templates" / "violation_protocols"
                if os.environ.get("PQM_DATA_DIR") else PACKAGED_TEMPLATE_DIR)
TEMPLATES = {
    "warning": TEMPLATE_DIR / "warning.docx",
    "decline_p49_1_2": TEMPLATE_DIR / "decline_p49_1_2.docx",
    "decline_p49_3": TEMPLATE_DIR / "decline_p49_3.docx",
}
TOKEN_RE = re.compile(r"\{\{\s*(.*?)\s*\}\}")


def ensure_runtime_templates() -> None:
    """Seed a persistent template directory without overwriting operator changes."""
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    for key, target in TEMPLATES.items():
        source = PACKAGED_TEMPLATE_DIR / target.name
        if not target.exists() and source.exists() and source.resolve() != target.resolve():
            shutil.copy2(source, target)


def template_metadata() -> list[dict[str, Any]]:
    ensure_runtime_templates()
    labels = {
        "warning": "Попередження за зверненням",
        "decline_p49_1_2": "Відмова за пп. 1–2 п. 49",
        "decline_p49_3": "Відмова за пп. 3 п. 49",
    }
    result = []
    for key, path in TEMPLATES.items():
        stat = path.stat() if path.exists() else None
        result.append({"key": key, "name": labels[key], "filename": path.name,
                       "exists": bool(stat),
                       "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(timespec="seconds") if stat else None})
    return result


def _template_tokens(path: Path) -> set[str]:
    document = Document(path)
    return {_norm(match.group(1)) for paragraph in _all_paragraphs(document)
            for match in TOKEN_RE.finditer(paragraph.text)}


def replace_runtime_template(key: str, source_path: str | Path) -> Path:
    """Validate and atomically activate a DOCX while retaining the previous version."""
    ensure_runtime_templates()
    target = TEMPLATES.get(key)
    source = Path(source_path)
    if not target or source.suffix.lower() != ".docx" or not zipfile.is_zipfile(source):
        raise ValueError("Потрібен коректний файл DOCX для вибраного шаблону")
    try:
        expected = _template_tokens(target)
        supplied = _template_tokens(source)
    except Exception as exc:
        raise ValueError("Не вдалося прочитати структуру DOCX") from exc
    missing = sorted(expected - supplied)
    if missing:
        raise ValueError("У DOCX відсутні обов’язкові маркери: " + ", ".join(missing))
    versions = TEMPLATE_DIR / "_versions"
    versions.mkdir(parents=True, exist_ok=True)
    if target.exists():
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        shutil.copy2(target, versions / f"{target.stem}_{stamp}{target.suffix}")
    temporary = target.with_suffix(".docx.new")
    shutil.copy2(source, temporary)
    temporary.replace(target)
    return target


def _norm(value: str) -> str:
    return re.sub(r"[\s_.–—/-]+", " ", (value or "").strip().lower())


def _all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in document.sections:
        yield from section.header.paragraphs
        yield from section.footer.paragraphs


def _replace_in_paragraph(paragraph, values: dict[str, str], highlighted: set[str]):
    text = paragraph.text
    matches = list(TOKEN_RE.finditer(text))
    if not matches:
        return
    result, cursor = [], 0
    used_highlight = False
    for match in matches:
        result.append(text[cursor:match.start()])
        key = _norm(match.group(1))
        result.append(str(values.get(key, "—") or "—"))
        used_highlight = used_highlight or key in highlighted
        cursor = match.end()
    result.append(text[cursor:])
    value = "".join(result)
    runs = paragraph.runs
    if runs:
        runs[0].text = value
        if used_highlight:
            runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _add_hyperlink(paragraph, label: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single")
    props.extend((color, underline)); run.append(props)
    node = OxmlElement("w:t"); node.text = label
    run.append(node); hyperlink.append(run); paragraph._p.append(hyperlink)


def _set_document_links(paragraph, documents: list[dict[str, Any]]):
    for run in paragraph.runs:
        run.text = ""
    seen = set()
    for document in documents:
        label = str(document.get("title") or document.get("name") or document.get("documentType") or "Документ").strip()
        url = str(document.get("url") or "").strip()
        marker = (label, url)
        if marker in seen or not url:
            continue
        if seen:
            paragraph.add_run("; ")
        _add_hyperlink(paragraph, label, url)
        seen.add(marker)
    if not seen:
        paragraph.add_run("не надано")


def _clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""


def _configure_customer_document_block(document, documents: list[dict[str, Any]]):
    """Keep the approved damaged-file wording only when an officer marked a file.

    The exact wording is already present in each approved DOCX.  We select the
    singular/plural paragraph rather than constructing or paraphrasing it.
    """
    unavailable = [item for item in documents if item.get("file_unavailable")]
    available = [item for item in documents if not item.get("file_unavailable")]
    for table in document.tables:
        for row in table.rows:
            if not any("докази порушення" in _norm(paragraph.text)
                       for cell in row.cells for paragraph in cell.paragraphs):
                continue
            target = row.cells[-1]
            paragraphs = target.paragraphs
            token = next((p for p in paragraphs if "докази порушення" in _norm(p.text)), None)
            singular = next((p for p in paragraphs if "додано файл" in _norm(p.text)
                             and "технічно пошкоджен" in _norm(p.text)), None)
            plural = next((p for p in paragraphs if "додані файли" in _norm(p.text)
                           and "технічно пошкоджен" in _norm(p.text)), None)
            if not token:
                return
            if singular:
                if len(unavailable) != 1:
                    singular._element.getparent().remove(singular._element)
            if plural:
                if len(unavailable) <= 1:
                    plural._element.getparent().remove(plural._element)
            # Remove empty separator paragraphs from the template.
            for paragraph in list(target.paragraphs):
                if paragraph is not token and not paragraph.text.strip():
                    paragraph._element.getparent().remove(paragraph._element)
            _set_document_links(token, unavailable if unavailable else documents)
            if unavailable and available:
                extra = target.add_paragraph()
                _set_document_links(extra, available)
            return


def _replace_justification(document, justification: str):
    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            if "обґрунтування рішення" not in row_text.lower():
                continue
            target = row.cells[-1]
            if len(row.cells) == 1:
                target = row.cells[0]
            paragraphs = target.paragraphs
            # Preserve the first paragraph formatting and remove all sample variants.
            first = paragraphs[0]
            for run in first.runs:
                run.text = ""
            first.add_run(justification)
            for paragraph in paragraphs[1:]:
                paragraph._element.getparent().remove(paragraph._element)
            return
    raise ValueError("У погодженому шаблоні не знайдено блок «Обґрунтування рішення»")


def _remove_optional_rows(document, flags: dict[str, bool]):
    labels = {
        "written_refusal": ("письмов", "відмов"),
        "contract": ("укладен", "договор"),
        "guarantee": ("забезпеченн", "договор"),
        "court": ("суд", "рішенн"),
    }
    for table in document.tables:
        for row in list(table.rows):
            text = " ".join(cell.text for cell in row.cells).lower()
            for key, needles in labels.items():
                if not flags.get(key, False) and all(needle in text for needle in needles):
                    row._element.getparent().remove(row._element)
                    break
    if not flags.get("civil_code", False):
        for paragraph in list(document.paragraphs):
            if "цивільн" in paragraph.text.lower() and "кодекс" in paragraph.text.lower():
                paragraph._element.getparent().remove(paragraph._element)


def build_violation_protocol_docx(
    protocol_type: str,
    output_path: str | Path,
    values: dict[str, str],
    justification: str,
    customer_documents: list[dict[str, Any]] | None = None,
    supplier_documents: list[dict[str, Any]] | None = None,
    flags: dict[str, bool] | None = None,
    highlighted_tokens: set[str] | None = None,
) -> Path:
    """Create a protocol by editing a copy and atomically publishing it.

    Writing directly to ``output_path`` is unsafe on Windows: an already-open
    protocol is exclusively locked by Word and ``copy2`` used to fail halfway
    through generation.  Build a complete sibling file first, then replace the
    public output in one operation.  A genuine external lock is reported as a
    controlled, user-facing conflict and never leaks the raw WinError.
    """
    ensure_runtime_templates()
    template = TEMPLATES.get(protocol_type)
    if not template or not template.exists():
        raise ValueError(f"Невідомий або відсутній шаблон протоколу: {protocol_type}")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_name(f".{output.stem}.{uuid.uuid4().hex}.tmp.docx")
    try:
        shutil.copy2(template, temporary)
        document = Document(temporary)
        _configure_customer_document_block(document, customer_documents or [])
        normalized_values = {_norm(key): str(value or "—") for key, value in values.items()}
        highlighted = {_norm(key) for key in (highlighted_tokens or set())}
        for paragraph in list(_all_paragraphs(document)):
            source = paragraph.text
            token_names = [_norm(m.group(1)) for m in TOKEN_RE.finditer(source)]
            if any("документ" in name and ("замов" in name or "скарж" in name) for name in token_names):
                _set_document_links(paragraph, customer_documents or [])
            elif any("документ" in name and ("постач" in name or "відпов" in name) for name in token_names):
                _set_document_links(paragraph, supplier_documents or [])
            else:
                _replace_in_paragraph(paragraph, normalized_values, highlighted)
        _replace_justification(document, justification)
        _remove_optional_rows(document, flags or {})
        unresolved = [p.text for p in _all_paragraphs(document) if "{{" in p.text or "}}" in p.text]
        if unresolved:
            raise ValueError("У DOCX залишилися незаповнені плейсхолдери")
        document.save(temporary)
        try:
            os.replace(temporary, output)
        except PermissionError as exc:
            raise PermissionError(
                f"Не вдалося оновити файл «{output.name}». "
                "Закрийте його у Microsoft Word або іншій програмі та повторіть формування."
            ) from exc
    finally:
        temporary.unlink(missing_ok=True)
    return output
