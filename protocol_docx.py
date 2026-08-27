"""DOCX protocol builder for PQM.

The document follows the approved Google Docs protocol structure while the
local MVP remains independent from Codex and Google OAuth.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _date(value: str, short: bool = False) -> str:
    raw = (value or "")[:10]
    try:
        parsed = datetime.strptime(raw, "%Y-%m-%d")
        return parsed.strftime("%d.%m.%y" if short else "%d.%m.%Y")
    except ValueError:
        return value or "—"


def _font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    return run


def _paragraph(
    doc, text="", *, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    before=0, after=0, size=11, first_line_indent=None,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if align is not None:
        p.alignment = align
    _font(p.add_run(text), size=size, bold=bold)
    return p


def _heading(doc, text):
    return _paragraph(doc, text, bold=True, before=8, after=0, size=12)


def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell(cell, text, *, bold=False, size=8, center=False, fill=None, allow_blank=False):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    value = "" if allow_blank and not text else str(text or "—")
    _font(p.add_run(value), size=size, bold=bold)
    if fill:
        _shade(cell, fill)


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def _set_table_grid(table, widths):
    """Apply the fixed Word table grid from the approved protocol sample."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _add_protocol_table(doc, rows, rejected=False):
    headers = [
        "№ з/п", "Назва учасника", "ЄДРПОУ / РНОКПП",
        "Ідентифікатор кваліфікації", "Код ДК 021:2015 та категорія товару",
        "Дата надходження заявки",
    ]
    if rejected:
        headers += ["Зауваження"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = (
        [566, 3103, 1142, 1981, 2383, 1109, 4412]
        if rejected else
        [567, 5524, 1275, 1984, 3682, 1664]
    )
    for index, header in enumerate(headers):
        _set_cell(table.rows[0].cells[index], header, bold=True, size=7.5, center=True, fill="D9EAF7")
    _repeat_header(table.rows[0])
    for number, item in enumerate(rows, 1):
        cells = table.add_row().cells
        values = [
            number,
            item.get("supplier_name"),
            item.get("supplier_code"),
            item.get("pretty_id"),
            f"{item.get('dk_code') or '—'} - {item.get('category_title') or '—'}",
            _date(item.get("date_published")),
        ]
        if rejected:
            values += [item.get("protocol_remarks") or "—"]
        for index, value in enumerate(values):
            _set_cell(cells[index], value, size=7.5, center=index in (0, 2, 5))
    _set_table_grid(table, widths)
    return table


def build_protocol_docx(payload: dict, output_path: Path) -> Path:
    items = payload["items"]
    admitted = [item for item in items if item.get("protocol_decision") == "admit"]
    rejected = [item for item in items if item.get("protocol_decision") == "reject"]
    number = payload["protocol_number"]
    protocol_date = payload["protocol_date"]
    date_from = payload["date_from"]
    date_to = payload["date_to"]
    officer = payload["officer"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    _paragraph(doc, f"ПРОТОКОЛ № {number}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    _paragraph(doc, "розгляду заявок учасників електронного каталогу Державної установи «Професійні закупівлі»", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _paragraph(doc, "уповноваженою особою Державної установи «Професійні закупівлі»", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell(meta.cell(0, 0), _date(protocol_date), size=11)
    _set_cell(meta.cell(0, 1), "Київ", size=11, center=False)
    _set_cell(meta.cell(1, 0), "", size=11, allow_blank=True)
    _set_cell(meta.cell(1, 1), "", size=11, allow_blank=True)
    _set_table_grid(meta, [5128, 5128])

    first_name = officer.split()[0].casefold() if officer else ""
    verb = "склала" if first_name.endswith(("а", "я")) else "склав"
    _paragraph(doc, f"Уповноважена особа електронного каталогу Державної установи «Професійні закупівлі» {officer} {verb} цей протокол (далі – Протокол) за результатами розгляду поданих заявок учасників, які були надані такими учасниками згідно з оголошеннями про проведення кваліфікації до електронного каталогу в електронній системі закупівель (далі – Каталог), оприлюдненими Державною установою «Професійні закупівлі» як адміністратором Каталогу.", first_line_indent=1.0)
    _paragraph(doc, "Розгляд поданих заявок здійснювався на підставі Закону України «Про публічні закупівлі», Порядку формування та використання електронного каталогу, затвердженого постановою Кабінету Міністрів України від 14.09.2020 № 822 (зі змінами) (далі – Порядок № 822), та Положення «Про порядок кваліфікації до електронного каталогу «Prozorro Market» в електронній системі закупівель Державною установою «Професійні закупівлі», затвердженого наказом від 12.12.2024 № 56 (зі змінами).", first_line_indent=1.0)

    _heading(doc, "ПОРЯДОК ДЕННИЙ:")
    _paragraph(doc, f"1. Розгляд заявок про приєднання до електронного каталогу, поданих учасниками за період {_date(date_from)} - {_date(date_to)}.")
    _paragraph(doc, "2. Прийняття рішення про включення учасників до переліку кваліфікованих постачальників за категорією товару, до якої учасник подав заявку.")
    _paragraph(doc, "3. Прийняття рішення про відхилення заявок учасників за результатами їх розгляду.")

    _heading(doc, "ВСТАНОВЛЕНО:")
    _heading(doc, "Питання 1:")
    _paragraph(doc, "Розглянуто заявки про приєднання до Каталогу згідно з переліком, наведеним у Таблиці 1 Додатка 1.")
    _heading(doc, "Питання 2:")
    _paragraph(doc, "За результатами розгляду заявок, зазначених у Таблиці 2 Додатка 1, не встановлено підстав для їх відхилення.")
    _heading(doc, "Питання 3:")
    _paragraph(doc, "Заявки, зазначені в Таблиці 3 Додатка 1, не відповідають вимогам Порядку № 822 та відповідних оголошень і підлягають відхиленню.")

    _heading(doc, "ВИРІШЕНО:")
    _paragraph(doc, f"1. Затвердити результати розгляду {len(items)} заявок, поданих за період {_date(date_from)} - {_date(date_to)}.")
    _paragraph(doc, f"2. Включити до переліку кваліфікованих постачальників учасників згідно з Таблицею 2 Додатка 1 ({len(admitted)} заявок).")
    _paragraph(doc, f"3. Відхилити заявки учасників згідно з Таблицею 3 Додатка 1 ({len(rejected)} заявок) із зазначенням виявлених невідповідностей.")
    _paragraph(doc, "4. Розмістити в електронній системі закупівель Протокол як повідомлення про результати розгляду заявок протягом одного робочого дня з дня прийняття рішення.")
    _paragraph(doc, "Декларування про відсутність конфлікту інтересів:", bold=True, size=9)
    _paragraph(doc, "[х] На виконання п. 4 наказу ДУ «Професійні закупівлі» від 15.04.2026 № 15 підтверджую відсутність приватного інтересу щодо зазначених юридичних осіб та фізичних осіб-підприємців, а також інших обставин, які можуть вплинути на об’єктивність або неупередженість прийняття рішення.", size=9)

    signature = doc.add_table(rows=1, cols=2)
    _set_cell(signature.cell(0, 0), "Уповноважена особа електронного каталогу", size=10)
    _set_cell(signature.cell(0, 1), officer, bold=True, size=10)
    _set_table_grid(signature, [5128, 5128])

    landscape = doc.add_section(WD_SECTION.NEW_PAGE)
    landscape.orientation = WD_ORIENT.LANDSCAPE
    landscape.page_width, landscape.page_height = landscape.page_height, landscape.page_width
    landscape.top_margin = Cm(1.2)
    landscape.bottom_margin = Cm(1.2)
    landscape.left_margin = Cm(1.0)
    landscape.right_margin = Cm(1.0)
    _paragraph(doc, f"Додаток № 1 до Протоколу від {_date(protocol_date)} № {number}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, after=6)
    _paragraph(doc, "Таблиця 1", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10, after=6)
    _add_protocol_table(doc, items)
    _paragraph(doc, "Таблиця 2", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, before=8, after=6, size=10)
    _add_protocol_table(doc, admitted)
    _paragraph(doc, "Таблиця 3", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, before=8, after=6, size=10)
    _add_protocol_table(doc, rejected, rejected=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
